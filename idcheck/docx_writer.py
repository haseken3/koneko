"""Word 出力モジュール — LLM 評価結果を Word ファイルに書き出す。"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

SECTION_SPACE_BEFORE = Pt(18)  # 大項目見出しの上余白
SECTION_SPACE_AFTER = Pt(6)    # 大項目見出しの下余白
ITEM_SPACE_BEFORE = Pt(12)     # 小項目見出しの上余白
ITEM_SPACE_AFTER = Pt(3)       # 小項目内の段落間

SCORE_LABEL = {
    -1: "要改善",
    0: "中立",
    1: "良好",
}

SCORE_COLOR = {
    -1: RGBColor(0xC0, 0x39, 0x2B),   # 赤
    0: RGBColor(0x7F, 0x8C, 0x8D),    # グレー
    1: RGBColor(0x27, 0xAE, 0x60),    # 緑
}


def _level_label(level) -> str:
    return f"L{level}"


def _add_meta_paragraph(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)


def build_docx(
    results: list[dict],
    pptx_filename: str,
    persona_key: str,
    model: str,
    levels_meta: list[dict] | None = None,
) -> BytesIO:
    """評価結果から Word ファイルを組み立てて BytesIO で返す。

    Args:
        results: evaluate_all の戻り値（list of {item, result}）
        pptx_filename: 元のPPTXファイル名
        persona_key: 評価ペルソナのキー
        model: 使用モデル
        levels_meta: checklist['levels'] のメタ情報（任意・header に使用）

    Returns:
        BytesIO: Word ファイルのバイナリ（st.download_button にそのまま渡せる）
    """
    doc = Document()

    # タイトル
    title = doc.add_heading(f"IDフィードバック（自動生成）", level=0)

    # メタ情報
    _add_meta_paragraph(doc, "対象資料", pptx_filename)
    _add_meta_paragraph(doc, "評価日時", datetime.now().strftime("%Y-%m-%d %H:%M"))

    # サマリ
    ok_results = [r for r in results if r["result"]["ok"]]
    fail_results = [r for r in results if not r["result"]["ok"]]

    score_counts = {-1: 0, 0: 0, 1: 0}
    for r in ok_results:
        score = r["result"]["evaluation"]["score"]
        score_counts[score] = score_counts.get(score, 0) + 1

    doc.add_heading("サマリ", level=1)
    summary_p = doc.add_paragraph()
    summary_p.add_run(
        f"評価項目: {len(results)}件（成功 {len(ok_results)} / 失敗 {len(fail_results)}）\n"
    )
    summary_p.add_run(
        f"スコア内訳: 良好 +1 = {score_counts[1]}件 / 中立 0 = {score_counts[0]}件 / "
        f"要改善 -1 = {score_counts[-1]}件\n"
    )
    if fail_results:
        summary_p.add_run(
            f"⚠ 評価失敗 {len(fail_results)} 件あり（API エラー等で評価できなかった項目）"
        ).font.color.rgb = SCORE_COLOR[-1]

    # 注意書き
    note_p = doc.add_paragraph()
    note_run = note_p.add_run(
        "本フィードバックは LLM による自動評価です。教員の最終確認・修正の参考としてご活用ください。"
    )
    note_run.italic = True
    note_run.font.size = Pt(9)

    # スコア別に項目をまとめる（要改善 → 中立 → 良好の順、各セクション内は id 順）
    doc.add_heading("評価結果", level=1)

    def _section_key(r):
        """ソート用キー: 失敗は最後尾、ok 項目はスコア -1/0/+1 の順、同 score 内は item_id 順。"""
        res = r["result"]
        if not res["ok"]:
            return (99, r["item"]["id"])
        return (res["evaluation"]["score"], r["item"]["id"])

    sorted_results = sorted(results, key=_section_key)

    by_score = {}
    for r in sorted_results:
        if not r["result"]["ok"]:
            key = "fail"
        else:
            key = r["result"]["evaluation"]["score"]
        by_score.setdefault(key, []).append(r)

    score_section_order = [
        (-1, "要改善が必要な項目"),
        (0, "判定が中立の項目"),
        (1, "良好な項目"),
        ("fail", "評価に失敗した項目"),
    ]

    for score_key, section_title in score_section_order:
        section_items = by_score.get(score_key, [])
        if not section_items:
            continue
        section_heading = doc.add_heading(
            f"{section_title}（{len(section_items)}件）", level=2
        )
        section_heading.paragraph_format.space_before = SECTION_SPACE_BEFORE
        section_heading.paragraph_format.space_after = SECTION_SPACE_AFTER

        for r in section_items:
            item = r["item"]
            res = r["result"]

            # 項目見出し
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = ITEM_SPACE_BEFORE
            p_title.paragraph_format.space_after = ITEM_SPACE_AFTER
            p_title.add_run(f"[{item['id']}] ").bold = True
            p_title.add_run(item["title"])

            if not res["ok"]:
                p_fail = doc.add_paragraph()
                p_fail.paragraph_format.space_after = ITEM_SPACE_AFTER
                run = p_fail.add_run(
                    f"⚠ 評価失敗: {res.get('error', '不明エラー')}"
                )
                run.font.color.rgb = SCORE_COLOR[-1]
                continue

            ev = res["evaluation"]
            score = ev["score"]

            # スコア表示
            p_score = doc.add_paragraph()
            p_score.paragraph_format.space_after = ITEM_SPACE_AFTER
            run = p_score.add_run(
                f"スコア: {score:+d} ({SCORE_LABEL.get(score, '?')})"
            )
            run.bold = True
            run.font.color.rgb = SCORE_COLOR.get(score, SCORE_COLOR[0])

            # rationale
            p_rat = doc.add_paragraph()
            p_rat.paragraph_format.space_after = ITEM_SPACE_AFTER
            p_rat.add_run("根拠: ").bold = True
            p_rat.add_run(ev["rationale"])

            # 該当スライド
            related = ev.get("related_slides", [])
            if related:
                p_rel = doc.add_paragraph()
                p_rel.paragraph_format.space_after = ITEM_SPACE_AFTER
                p_rel.add_run("該当スライド: ").bold = True
                p_rel.add_run(", ".join(f"S{n}" for n in related))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
